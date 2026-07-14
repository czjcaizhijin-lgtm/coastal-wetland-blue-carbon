from __future__ import annotations

from pathlib import Path
import json
import math
import sys
from typing import Iterable

import joblib
import numpy as np
import pandas as pd

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BASE = Path(r"D:\PyCharm2020(64bit)\pythonProject")
OUT_DIR = BASE / "不确定性空间制图" / "Uncertainty"
DATA_DIR = BASE / "处理" / "筛选后的反演点"
MODEL_PATH = OUT_DIR / "ensemble_30_models_NoLog.pkl"

FEATURE_COLS = [
    "ndvi",
    "BIO1_Annual_Temp",
    "Salinity_psu",
    "tsm_p",
    "BIO14_Precip_Dry_Month",
    "Vapr_kPa",
    "sst_k",
    "tidal_range_m",
    "soil_moisture",
    "bdod_1m",
    "clay_1m",
    "LSWI",
    "Night_Light",
    "BIO4_Temp_Seasonality",
    "BIO6_Min_Temp_Coldest",
]

CONTINENT_FILES = [
    ("Africa", DATA_DIR / "Africa_Standardized_Results.csv"),
    ("Asia", DATA_DIR / "Asia_Final_Standardized.csv"),
    ("Europe", DATA_DIR / "Europe_Standardized_Results.csv"),
    ("North America", DATA_DIR / "North_America_Final_Standardized.csv"),
    ("Oceania", DATA_DIR / "Oceania_Standardized_Results.csv"),
    ("South America", DATA_DIR / "South_America_Final_Standardized.csv"),
]

CHUNK_SIZE = 250_000


def valid_mask(df: pd.DataFrame) -> pd.Series:
    return (
        (df["ndvi"] >= 0.01)
        & (df["ndvi"] <= 0.25)
        & (df["Night_Light"] < 10)
    )


def area_weights(latitude: np.ndarray) -> np.ndarray:
    # Source grids use regular lon/lat spacing; cell area is proportional to cos(latitude).
    weights = np.cos(np.deg2rad(latitude.astype(float)))
    return np.clip(weights, 0.0, None)


def weighted_quantile(values: np.ndarray, weights: np.ndarray, quantiles: Iterable[float]) -> list[float]:
    values = np.asarray(values, dtype=float)
    weights = np.asarray(weights, dtype=float)
    mask = np.isfinite(values) & np.isfinite(weights) & (weights > 0)
    values = values[mask]
    weights = weights[mask]
    if len(values) == 0:
        return [math.nan for _ in quantiles]
    order = np.argsort(values)
    values = values[order]
    weights = weights[order]
    cumulative = np.cumsum(weights)
    cutoff = np.asarray(list(quantiles), dtype=float) * cumulative[-1]
    return np.interp(cutoff, cumulative, values).tolist()


def iter_valid_chunks(path: Path) -> Iterable[pd.DataFrame]:
    usecols = list(dict.fromkeys(FEATURE_COLS + ["longitude", "latitude"]))
    for chunk in pd.read_csv(path, usecols=usecols, chunksize=CHUNK_SIZE):
        filtered = chunk.loc[valid_mask(chunk)].copy()
        if not filtered.empty:
            yield filtered


def relative_error(models: list, df: pd.DataFrame) -> np.ndarray:
    x = df[FEATURE_COLS]
    predictions = [model.predict(x) for model in models]
    matrix = np.vstack(predictions)
    raw_sd = np.std(matrix, axis=0)
    raw_mean = np.mean(matrix, axis=0)
    return raw_sd / (raw_mean + 0.05)


def nui_from_relative_error(relative: np.ndarray, min_err: float, max_err: float) -> np.ndarray:
    nui = 0.05 + 0.45 * (relative - min_err) / (max_err - min_err)
    return np.clip(nui, 0.0, 0.50)


def cache_path(continent: str) -> Path:
    safe_name = continent.replace(" ", "_")
    return OUT_DIR / f"Supplementary_Table_9_intermediate_{safe_name}.npz"


def load_or_compute_continent_arrays(models: list, continent: str, path: Path) -> tuple[np.ndarray, np.ndarray]:
    target = cache_path(continent)
    if target.exists():
        print(f"[cache] {continent}: loading {target}", flush=True)
        loaded = np.load(target)
        return loaded["relative"], loaded["weights"]

    relative_parts = []
    weight_parts = []
    valid_count = 0
    print(f"[predict] {continent}: {path.name}", flush=True)
    for chunk in iter_valid_chunks(path):
        rel = relative_error(models, chunk)
        weights = area_weights(chunk["latitude"].to_numpy(float))
        relative_parts.append(rel.astype(np.float32))
        weight_parts.append(weights.astype(np.float32))
        valid_count += len(rel)
        print(f"  valid pixels processed: {valid_count:,}", flush=True)

    relative = np.concatenate(relative_parts)
    weights = np.concatenate(weight_parts)
    np.savez(target, relative=relative, weights=weights)
    print(f"  cached: {target}", flush=True)
    return relative, weights


def collect_relative_errors(models: list) -> tuple[dict[str, tuple[np.ndarray, np.ndarray]], dict[str, int]]:
    arrays_by_continent = {}
    all_relative_parts = []
    counts = {}
    for continent, path in CONTINENT_FILES:
        relative, weights = load_or_compute_continent_arrays(models, continent, path)
        arrays_by_continent[continent] = (relative, weights)
        all_relative_parts.append(relative)
        counts[continent] = int(len(relative))
        print(f"  valid pixels total: {len(relative):,}", flush=True)
    return arrays_by_continent, counts, np.concatenate(all_relative_parts)


def summarise_by_continent(
    arrays_by_continent: dict[str, tuple[np.ndarray, np.ndarray]],
    min_err: float,
    max_err: float,
    high_threshold: float,
) -> pd.DataFrame:
    rows = []
    for continent, (relative, weights) in arrays_by_continent.items():
        values = nui_from_relative_error(relative, min_err, max_err)
        q25, median, q75 = weighted_quantile(values, weights, [0.25, 0.50, 0.75])
        high_area = float(weights[values >= high_threshold].sum() / weights.sum() * 100.0)
        mean_nui = float(np.average(values, weights=weights))

        rows.append(
            {
                "Region": continent,
                "Mean NUI": mean_nui,
                "Median NUI": median,
                "NUI 25th percentile": q25,
                "NUI 75th percentile": q75,
                "High-uncertainty area (%)": high_area,
                "Valid pixels": int(len(values)),
            }
        )
    return pd.DataFrame(rows)


def write_xlsx(df: pd.DataFrame, path: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.chart import BarChart, LineChart, Reference
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except Exception as exc:
        print(f"[warning] openpyxl unavailable, skipped xlsx export: {exc}")
        return

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Supplementary Table 9"
    headers = list(df.columns)
    sheet.append(headers)
    for row in df.itertuples(index=False):
        sheet.append(list(row))

    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row in sheet.iter_rows(min_row=2):
        for idx, cell in enumerate(row, start=1):
            cell.alignment = Alignment(horizontal="right" if idx > 1 else "left", vertical="top", wrap_text=True)

    widths = [20, 14, 14, 20, 20, 24, 14]
    for idx, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(idx)].width = width

    for col_idx in range(2, 6):
        for row_idx in range(2, sheet.max_row + 1):
            sheet.cell(row=row_idx, column=col_idx).number_format = "0.000"
    for row_idx in range(2, sheet.max_row + 1):
        sheet.cell(row=row_idx, column=6).number_format = "0.0"
        sheet.cell(row=row_idx, column=7).number_format = "#,##0"

    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions

    chart_sheet = workbook.create_sheet("Visualization")
    chart_sheet["A1"] = "Regional uncertainty statistics of predicted deep SOC stocks"
    chart_sheet["A1"].font = Font(bold=True, size=12)
    chart_sheet["A3"] = "Region"
    chart_sheet["B3"] = "Mean NUI"
    chart_sheet["C3"] = "High-uncertainty area (%)"
    for row_idx, row in enumerate(df.itertuples(index=False), start=4):
        chart_sheet.cell(row=row_idx, column=1, value=row[0])
        chart_sheet.cell(row=row_idx, column=2, value=row[1])
        chart_sheet.cell(row=row_idx, column=3, value=row[5])
    for cell in chart_sheet[3]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    for row_idx in range(4, 4 + len(df)):
        chart_sheet.cell(row=row_idx, column=2).number_format = "0.000"
        chart_sheet.cell(row=row_idx, column=3).number_format = "0.0"
    chart_sheet.column_dimensions["A"].width = 20
    chart_sheet.column_dimensions["B"].width = 14
    chart_sheet.column_dimensions["C"].width = 24

    categories = Reference(chart_sheet, min_col=1, min_row=4, max_row=3 + len(df))
    mean_data = Reference(chart_sheet, min_col=2, min_row=3, max_row=3 + len(df))
    high_data = Reference(chart_sheet, min_col=3, min_row=3, max_row=3 + len(df))

    bar = BarChart()
    bar.title = "Mean NUI by continent"
    bar.y_axis.title = "Mean NUI"
    bar.x_axis.title = "Region"
    bar.add_data(mean_data, titles_from_data=True)
    bar.set_categories(categories)
    bar.height = 8
    bar.width = 16
    chart_sheet.add_chart(bar, "E3")

    line = LineChart()
    line.title = "High-uncertainty area by continent"
    line.y_axis.title = "Area (%)"
    line.x_axis.title = "Region"
    line.add_data(high_data, titles_from_data=True)
    line.set_categories(categories)
    line.height = 8
    line.width = 16
    chart_sheet.add_chart(line, "E21")

    workbook.save(path)


def write_png(df: pd.DataFrame, path: Path) -> None:
    try:
        import matplotlib.pyplot as plt
    except Exception as exc:
        print(f"[warning] matplotlib unavailable, skipped png export: {exc}")
        return

    regions = df["Region"].tolist()
    mean_nui = df["Mean NUI"].to_numpy(float)
    high_area = df["High-uncertainty area (%)"].to_numpy(float)
    x = np.arange(len(regions))

    plt.rcParams.update({"font.size": 10, "font.family": "Arial"})
    fig, ax1 = plt.subplots(figsize=(9.2, 4.8), dpi=300)
    bars = ax1.bar(x - 0.18, mean_nui, width=0.36, color="#4E79A7", label="Mean NUI")
    ax1.set_ylabel("Mean NUI")
    ax1.set_ylim(0, max(0.35, float(mean_nui.max()) * 1.25))
    ax1.set_xticks(x)
    ax1.set_xticklabels(regions, rotation=25, ha="right")

    ax2 = ax1.twinx()
    ax2.plot(x + 0.18, high_area, color="#E15759", marker="o", linewidth=2, label="High-uncertainty area")
    ax2.set_ylabel("High-uncertainty area (%)")
    ax2.set_ylim(0, max(50, float(high_area.max()) * 1.25))

    ax1.grid(axis="y", linestyle="--", alpha=0.25)
    ax1.set_title("Supplementary Table 9. Regional uncertainty statistics")
    handles = [bars, ax2.lines[0]]
    labels = ["Mean NUI", "High-uncertainty area (%)"]
    ax1.legend(handles, labels, loc="upper right", frameon=False)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def write_docx(df: pd.DataFrame, path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except Exception as exc:
        print(f"[warning] python-docx unavailable, skipped docx export: {exc}")
        return

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Supplementary Table 9. Regional uncertainty statistics of predicted deep SOC stocks.")
    run.bold = True
    run.font.size = Pt(10.5)

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, header in enumerate(df.columns):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(8.5)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(df.columns):
            value = row[column]
            if column in {"Mean NUI", "Median NUI", "NUI 25th percentile", "NUI 75th percentile"}:
                text = f"{value:.3f}"
            elif column == "High-uncertainty area (%)":
                text = f"{value:.1f}"
            elif column == "Valid pixels":
                text = f"{int(value):,}"
            else:
                text = str(value)
            run = cells[idx].paragraphs[0].add_run(text)
            run.font.size = Pt(8)
    document.save(path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Loading ensemble models: {MODEL_PATH}", flush=True)
    models = joblib.load(MODEL_PATH)
    print(f"Loaded models: {len(models)}", flush=True)

    arrays_by_continent, valid_counts, relative_all = collect_relative_errors(models)
    min_err, max_err = np.nanpercentile(relative_all, [2, 98])
    global_nui = nui_from_relative_error(relative_all, float(min_err), float(max_err))
    high_threshold = float(np.nanpercentile(global_nui, 75))
    print(f"NUI scaling relative-error p2={min_err:.6f}, p98={max_err:.6f}", flush=True)
    print(f"Global high-uncertainty threshold: NUI >= {high_threshold:.6f}", flush=True)

    df = summarise_by_continent(arrays_by_continent, float(min_err), float(max_err), high_threshold)
    df = df[
        [
            "Region",
            "Mean NUI",
            "Median NUI",
            "NUI 25th percentile",
            "NUI 75th percentile",
            "High-uncertainty area (%)",
            "Valid pixels",
        ]
    ]

    csv_path = OUT_DIR / "Supplementary_Table_9_regional_uncertainty_statistics_by_continent.csv"
    xlsx_path = OUT_DIR / "Supplementary_Table_9_regional_uncertainty_statistics_by_continent.xlsx"
    docx_path = OUT_DIR / "Supplementary_Table_9_regional_uncertainty_statistics_by_continent.docx"
    png_path = OUT_DIR / "Supplementary_Table_9_regional_uncertainty_statistics_by_continent.png"
    notes_path = OUT_DIR / "Supplementary_Table_9_source_notes.json"

    df.to_csv(csv_path, index=False, encoding="utf-8-sig")
    write_xlsx(df, xlsx_path)
    write_png(df, png_path)
    write_docx(df, docx_path)

    notes = {
        "title": "Supplementary Table 9. Regional uncertainty statistics of predicted deep SOC stocks.",
        "region_definition": "Continents were defined by the six continent-specific standardized prediction CSV files in the source data folder. The mixed Asia_Europe_Final_Standardized.csv file was excluded to avoid double-counting Asia and Europe.",
        "nui_method": "NUI was reproduced from the saved 30-model ensemble following the mapping used in the uncertainty scripts: relative error = ensemble SD / (ensemble mean + 0.05); the 2nd and 98th percentiles of all valid continent pixels define the min-max scaling to 0.05-0.50, followed by clipping to 0.00-0.50.",
        "valid_pixel_filter": "ndvi >= 0.01, ndvi <= 0.25, and Night_Light < 10.",
        "area_weighting": "Mean NUI, quantiles, and high-uncertainty percentages were area-weighted using cos(latitude), appropriate for regular longitude-latitude grids.",
        "high_uncertainty_definition": f"High-uncertainty area is the area-weighted percentage of valid pixels with NUI >= global 75th percentile ({high_threshold:.6f}).",
        "model_file": str(MODEL_PATH),
        "source_files": {continent: str(path) for continent, path in CONTINENT_FILES},
        "valid_pixel_counts": valid_counts,
    }
    notes_path.write_text(json.dumps(notes, ensure_ascii=False, indent=2), encoding="utf-8")

    print(csv_path)
    print(xlsx_path)
    print(png_path)
    if docx_path.exists():
        print(docx_path)
    print(notes_path)
    print(df.to_string(index=False))


if __name__ == "__main__":
    main()
